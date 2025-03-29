import os
from typing import Dict, Any, List, Optional
from docx import Document
from prompts_old import ERROR_MESSAGES


def analyze_document(doc_path: str) -> str:
    if not os.path.exists(doc_path):
        raise FileNotFoundError(ERROR_MESSAGES['file_not_found'].format(path=doc_path))

    if not doc_path.endswith('.docx'):
        raise ValueError(ERROR_MESSAGES['invalid_format'].format(path=doc_path))

    try:
        doc: Document = Document(doc_path)
        content: List[str] = []
        last_paragraph: Optional[str] = None  # Для хранения последнего текста перед таблицей

    # Проходим по всем элементам документа (абзацы и таблицы)
        for element in doc.element.body:
            if element.tag.endswith("p"):  # Абзац
                para = next((p for p in doc.paragraphs if p._element == element), None)
                if para and para.text.strip():
                    last_paragraph = para.text  # Запоминаем текст перед таблицей
                    content.append(para.text)
            elif element.tag.endswith("tbl"):  # Таблица
                table = next((t for t in doc.tables if t._element == element), None)
                if table:
                    table_title = last_paragraph if last_paragraph else "Без названия"
                    table_str = "\n".join(
                        " | ".join(cell.text.strip() for cell in row.cells) for row in table.rows
                    )
                    content.append(f"\n[Таблица: {table_title}]\n{table_str}\n")
                    last_paragraph = None  # Сбрасываем, чтобы не использовать один и тот же заголовок дважды

        return "\n".join(content)

    except Exception as e:
        raise Exception(ERROR_MESSAGES['doc_read_error'].format(error=str(e)))
